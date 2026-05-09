"""
Question Bank Converter  ·  Word (.docx) → ONE Master Excel (.xlsx)
====================================================================
Handles ALL sections including Section E (Case-Based / Case Study).
Works even when ALL paragraphs use Normal style (no Heading styles).

HOW TO RUN:
  1. Install dependencies once:
       pip install python-docx openpyxl
  2. Run:
       python docx_to_excel_converter.py
  3. Enter INPUT folder and OUTPUT folder when prompted.

OUTPUT: QuestionBank_Master.xlsx
  • One worksheet per CLASS
  • Columns: Class | Subject | Chapter | Chapter Name | Case Text |
             Q No | Question | Question Type | Option 1-4 |
             Correct Answer | Marks | Difficulty Level
  • AutoFilter on every column
"""

import sys, os, re
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter


# ══════════════════════════════════════════════════════════════════════
#  HELPERS
# ══════════════════════════════════════════════════════════════════════

def extract_difficulty(text):
    m = re.search(r'\((Low|Mid|High|Advance)\s+Level\)', text, re.IGNORECASE)
    if m:
        return {'low':'L','mid':'M','high':'H','advance':'A'}.get(m.group(1).lower(),'')
    m = re.search(r'\[(L|M|A)\]', text)
    return m.group(1) if m else ''

def clean_text(text):
    text = re.sub(r'\((Low|Mid|High|Advance)\s+Level\)', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\[(L|M|A)\]', '', text)
    # Remove trailing punctuation noise left by removing difficulty tag
    return text.strip().rstrip(',').strip()

def clean_option(text):
    return re.sub(r'^\s*\(?[a-dA-D]\)?\s*', '', text).strip()

def extract_inline_answer(lines):
    """Return (answer_letter_or_text, index_of_answer_line_or_-1)."""
    for i, line in enumerate(lines):
        if re.match(r'^Answer\s*[:=]', line, re.IGNORECASE):
            raw = re.sub(r'^Answer\s*[:=]\s*', '', line, flags=re.IGNORECASE).strip()
            m = re.match(r'^\(?([a-dA-D])\)?', raw)
            return (m.group(1).lower() if m else raw), i
    return '', -1

# ── Subject detection ─────────────────────────────────────────────────
KNOWN_SUBJECTS = [
    'Social Science','Social Studies','Environmental Science',
    'Computer Science','Information Technology','Political Science',
    'Business Studies','Physical Education',
    'Mathematics','Science','Physics','Chemistry','Biology',
    'English','Hindi','Sanskrit','Urdu','Marathi','Tamil',
    'History','Geography','Civics','Economics','Accountancy','Commerce',
    'Computer','EVS','Maths','Math','SST','IT',
]
_SNORM = {
    'maths':'Mathematics','math':'Mathematics','sst':'Social Science',
    'social studies':'Social Science','evs':'Environmental Science',
    'it':'Information Technology','computer':'Computer Science',
}
def _norm_subj(raw): return _SNORM.get(raw.lower().strip(), raw.strip().title())

def extract_subject(doc, filepath):
    texts = [p.text.strip() for p in doc.paragraphs[:15] if p.text.strip()]
    texts.append(os.path.splitext(os.path.basename(filepath))[0])
    for text in texts:
        for subj in KNOWN_SUBJECTS:
            if re.search(r'\b'+re.escape(subj)+r'\b', text, re.IGNORECASE):
                return _norm_subj(subj)
    return 'General'

def sanitize_sheet(name):
    return re.sub(r'[/\\?*\[\]:]','',name)[:31].strip() or 'Sheet'


# ══════════════════════════════════════════════════════════════════════
#  PARSER
# ══════════════════════════════════════════════════════════════════════

def parse_document(docx_path):
    doc = Document(docx_path)
    subject = extract_subject(doc, docx_path)

    # Detect class name
    class_name = 'Unknown Class'
    for para in doc.paragraphs[:10]:
        m = re.search(r'Class\s*(\d+)', para.text, re.IGNORECASE)
        if m:
            class_name = f'Class {m.group(1)}'
            break

    # ── Mutable state ─────────────────────────────────────────────────
    sec          = None          # current section: 'A','B','C','D','E', or None
    chapter      = 'Chapter 1'
    ch_name      = ''
    grab_ch_name = False         # True = next non-meta line is the chapter title

    # Section-E state
    in_case      = False
    case_num     = 0
    case_text    = ''
    in_q_block   = False         # True after "Answer the following..." / "Questions:"
    sub_idx      = 0             # counter for Pattern-B sub-questions

    # Per-section counters (reset every new chapter)
    ctr = {k:1 for k in ['MCQ','TF','FIB','VSAQ','SAQ','LAQ']}

    all_rows = []

    def add(q_type, q_no, question, o1, o2, o3, o4, answer, marks, diff, case_t=''):
        all_rows.append({
            'Class':class_name,  'Subject':subject,
            'Chapter':chapter,   'Chapter Name':ch_name,
            'Case Text':case_t,
            'Q No':q_no,         'Question':question,
            'Question Type':q_type,
            'Option 1':o1,'Option 2':o2,'Option 3':o3,'Option 4':o4,
            'Correct Answer':answer,
            'Marks':marks,       'Difficulty Level':diff,
        })

    # ── Walk every paragraph ──────────────────────────────────────────
    for para in doc.paragraphs:
        raw  = para.text          # original (may have embedded \n)
        text = raw.strip()
        if not text:
            continue

        # ════════════════════════════════════════════════════════════
        #  STEP 1 – Answer Key line  (MUST come before section detect)
        #  "Answer Key (Chapter 1 Section A)" contains "Section A"
        #  so we must intercept it before the section-detection step.
        # ════════════════════════════════════════════════════════════
        if re.match(r'^Answer\s+Key', text, re.IGNORECASE):
            sec = 'KEY'
            continue

        # ════════════════════════════════════════════════════════════
        #  STEP 2 – Chapter detection (runs even inside KEY section
        #           so we can resume parsing the next chapter)
        # ════════════════════════════════════════════════════════════
        # Pattern A: "Chapter 1: Title"  or  "... – Chapter 3: Title"
        m = re.search(r'\bChapter\s+(\d+)\s*[:\-–]\s*(.+)', text)
        if m:
            chapter  = f'Chapter {m.group(1)}'
            ch_name  = clean_text(m.group(2).strip())
            ctr      = {k:1 for k in ctr}
            sec      = None
            in_case  = False; in_q_block = False; case_text = ''
            grab_ch_name = False
            continue

        # Pattern B: line ends with "Chapter N"  (title on next line)
        m = re.match(r'^.*\bChapter\s+(\d+)\s*$', text)
        if m:
            chapter  = f'Chapter {m.group(1)}'
            ch_name  = ''
            ctr      = {k:1 for k in ctr}
            sec      = None
            in_case  = False; in_q_block = False; case_text = ''
            grab_ch_name = True
            continue

        # Capture chapter title from the line following Pattern-B header
        if grab_ch_name:
            grab_ch_name = False
            # Skip lines that look like meta / section headers
            if not re.match(r'^(Section\s+[A-E]|\(|\d+\s+Q|Comprehensive)',
                            text, re.IGNORECASE):
                ch_name = clean_text(text)
                continue
            # else fall through and let the line be processed normally

        # ════════════════════════════════════════════════════════════
        #  STEP 3 – Skip everything in the Answer Key section
        # ════════════════════════════════════════════════════════════
        if sec == 'KEY':
            continue

        # ════════════════════════════════════════════════════════════
        #  STEP 4 – Section header detection
        #  Only AFTER we've already handled "Answer Key (…Section A…)"
        # ════════════════════════════════════════════════════════════
        m = re.search(r'\bSection\s+([A-E])\b', text, re.IGNORECASE)
        if m:
            sec     = m.group(1).upper()
            in_case = False
            in_q_block = False
            continue

        # ════════════════════════════════════════════════════════════
        #  STEP 5 – Skip meta / boilerplate lines
        # ════════════════════════════════════════════════════════════
        if re.match(r'^\(\d+\s+Marks?\s+Each\)', text, re.IGNORECASE): continue
        if re.match(r'^\(\d+\s+Mark', text, re.IGNORECASE):            continue
        if re.match(r'^\d+\s+Questions?\s*$', text, re.IGNORECASE):    continue
        if re.match(r'^Comprehensive\s+Question\s+Bank', text, re.IGNORECASE): continue
        if re.match(r'^Question\s+Bank\s*[–\-]', text, re.IGNORECASE): continue

        # No section yet → skip
        if sec is None:
            continue

        # ════════════════════════════════════════════════════════════
        #  SECTION A  –  Objective Questions
        # ════════════════════════════════════════════════════════════
        if sec == 'A':
            lines = [l.strip() for l in raw.split('\n') if l.strip()]
            if not lines: continue

            q_line = lines[0]
            diff   = extract_difficulty(raw)

            # Collect option lines  (a) …  (b) …  (c) …  (d) …
            opts   = [l for l in lines[1:]
                      if re.match(r'^\(?[a-dA-D]\)', l) and
                         not re.match(r'^Answer\s*[:=]', l, re.IGNORECASE)]

            # Extract inline answer (newer chapters)
            answer_val, _ = extract_inline_answer(lines)

            q_clean = clean_text(q_line)

            # ── Detect True/False ────────────────────────────────────
            is_tf = bool(re.search(r'true\s*or\s*false|t\s*/\s*f', q_line, re.IGNORECASE))
            if not is_tf:
                ans_l = next((l for l in lines
                              if re.match(r'^Answer\s*[:=]', l, re.IGNORECASE)), '')
                if ans_l:
                    val = re.sub(r'^Answer\s*[:=]\s*', '', ans_l, flags=re.IGNORECASE).strip().lower()
                    if re.match(r'^(true|false|a\)\s*true|b\)\s*false)', val):
                        is_tf = True
            if not is_tf and len(opts) == 2:
                op_set = {clean_option(o).lower() for o in opts}
                if op_set == {'true', 'false'}:
                    is_tf = True

            # ── Detect Fill-in-Blank ─────────────────────────────────
            is_fib = bool(re.search(r'fill\s+in|_{3,}', q_line, re.IGNORECASE))

            # ── Classify & store ─────────────────────────────────────
            if is_tf:
                q_clean = re.sub(r'^True\s+or\s+False\s*[:–\-]?\s*', '', q_clean,
                                 flags=re.IGNORECASE).strip()
                add('True False', ctr['TF'], q_clean,
                    'True','False','','', answer_val, 1, diff)
                ctr['TF'] += 1
            elif is_fib:
                add('Fill in Blank', ctr['FIB'], q_clean,
                    '','','','', answer_val, 1, diff)
                ctr['FIB'] += 1
            else:
                o = ['','','','']
                for i,op in enumerate(opts[:4]): o[i] = clean_option(op)
                add('MCQ', ctr['MCQ'], q_clean,
                    o[0],o[1],o[2],o[3], answer_val, 1, diff)
                ctr['MCQ'] += 1

        # ════════════════════════════════════════════════════════════
        #  SECTION B – Very Short Answer Questions
        # ════════════════════════════════════════════════════════════
        elif sec == 'B':
            add('VSAQ', ctr['VSAQ'], clean_text(text),
                '','','','', '', 2, extract_difficulty(text))
            ctr['VSAQ'] += 1

        # ════════════════════════════════════════════════════════════
        #  SECTION C – Short Answer Questions
        # ════════════════════════════════════════════════════════════
        elif sec == 'C':
            add('SAQ', ctr['SAQ'], clean_text(text),
                '','','','', '', 3, extract_difficulty(text))
            ctr['SAQ'] += 1

        # ════════════════════════════════════════════════════════════
        #  SECTION D – Long Answer Questions
        # ════════════════════════════════════════════════════════════
        elif sec == 'D':
            add('LAQ', ctr['LAQ'], clean_text(text),
                '','','','', '', 5, extract_difficulty(text))
            ctr['LAQ'] += 1

        # ════════════════════════════════════════════════════════════
        #  SECTION E – Case-Based / Case Study Questions
        # ════════════════════════════════════════════════════════════
        elif sec == 'E':

            # ── Case header: "Case-Based Question 1" / "Case Study 2" ─
            if re.match(r'^Case[-\s]*Based\s+Question\s*\d+|^Case\s+Study\s*\d+',
                        text, re.IGNORECASE):
                m_n     = re.search(r'\d+', text)
                case_num   = int(m_n.group()) if m_n else (case_num + 1)
                case_text  = ''
                in_case    = True
                in_q_block = False
                sub_idx    = 0
                continue          # ← do NOT treat header as a question

            if not in_case:
                continue

            # ── "Answer the following questions:" / "Questions:" ──────
            # Using a broad regex to catch all variants:
            #   "Answer the following questions:"
            #   "Answer the following questions"
            #   "Questions:"
            #   "Questions"
            if re.match(r'^Answer\s+the\s+following|^Questions?\s*[:.]?\s*$',
                        text, re.IGNORECASE):
                in_q_block = True
                continue          # ← do NOT treat trigger as a question

            # ── Case scenario paragraph (before questions block) ──────
            if not in_q_block:
                case_text = clean_text(text)
                continue          # ← do NOT treat scenario as a question

            # ── Sub-questions ─────────────────────────────────────────
            lines = [l.strip() for l in raw.split('\n') if l.strip()]
            if not lines:
                continue

            first = lines[0]

            # PATTERN A: All sub-questions in ONE paragraph
            #   "a) What problem…\nb) What did…\nc) …"
            if re.match(r'^[a-f]\)', first):
                for line in lines:
                    m_s = re.match(r'^([a-f])\)\s*(.+)', line)
                    if m_s:
                        sub_letter = m_s.group(1)
                        sub_q      = clean_text(m_s.group(2))
                        diff       = extract_difficulty(line)
                        add('Case Based',
                            f'Case {case_num}({sub_letter})',
                            sub_q, '','','','', '', 1, diff, case_text)
                continue

            # PATTERN B: Each sub-question is its own paragraph
            sub_idx += 1
            q_sub = clean_text(text)
            diff  = extract_difficulty(text)
            add('Case Based',
                f'Case {case_num}({sub_idx})',
                q_sub, '','','','', '', 1, diff, case_text)

    return class_name, subject, all_rows


# ══════════════════════════════════════════════════════════════════════
#  EXCEL WRITER
# ══════════════════════════════════════════════════════════════════════

COLUMNS    = ['Class','Subject','Chapter','Chapter Name','Case Text',
              'Q No','Question','Question Type',
              'Option 1','Option 2','Option 3','Option 4',
              'Correct Answer','Marks','Difficulty Level']
COL_WIDTHS = [10, 14, 10, 22, 40, 12, 46, 14, 14, 14, 14, 14, 14, 7, 14]
CENTER_COL = {'Q No','Marks','Difficulty Level','Question Type','Class',
              'Chapter','Subject','Option 1','Option 2','Option 3',
              'Option 4','Correct Answer'}

_HF  = PatternFill('solid', start_color='4472C4')
_HFN = Font(bold=True, color='FFFFFF', name='Arial', size=10)
_CF  = Font(name='Arial', size=10)
_WA  = Alignment(wrap_text=True, vertical='top')
_CA  = Alignment(horizontal='center', vertical='top', wrap_text=True)
_T   = Side(border_style='thin', color='CCCCCC')
_BD  = Border(left=_T, right=_T, top=_T, bottom=_T)
_AF  = PatternFill('solid', start_color='EEF2FA')
_TF  = PatternFill('solid', start_color='D9E1F2')


def write_excel(classes_data, output_path):
    wb = Workbook()
    wb.remove(wb.active)

    for cls in sorted(classes_data):
        rows = classes_data[cls]
        if not rows:
            continue
        ws = wb.create_sheet(title=sanitize_sheet(cls))

        # Title row
        ws.append(['Question Bank – ' + cls] + ['']*(len(COLUMNS)-1))
        ws.merge_cells(start_row=1, start_column=1,
                       end_row=1, end_column=len(COLUMNS))
        c = ws.cell(1,1)
        c.font      = Font(bold=True, name='Arial', size=11, color='1F3864')
        c.fill      = _TF
        c.alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[1].height = 22

        # Header row
        ws.append(COLUMNS)
        for ci in range(1, len(COLUMNS)+1):
            cell = ws.cell(2, ci)
            cell.font = _HFN; cell.fill = _HF
            cell.alignment = _CA; cell.border = _BD
        ws.row_dimensions[2].height = 18

        # Data rows
        for rd in rows:
            ws.append([rd.get(c,'') for c in COLUMNS])
            rn = ws.max_row
            ws.row_dimensions[rn].height = 32
            for ci, col in enumerate(COLUMNS, 1):
                cell = ws.cell(rn, ci)
                cell.font   = _CF
                cell.border = _BD
                cell.alignment = _CA if col in CENTER_COL else _WA
                if rn % 2 == 0:
                    cell.fill = _AF

        # Column widths, AutoFilter, Freeze
        for ci, w in enumerate(COL_WIDTHS, 1):
            ws.column_dimensions[get_column_letter(ci)].width = w
        ws.auto_filter.ref = f'A2:{get_column_letter(len(COLUMNS))}{ws.max_row}'
        ws.freeze_panes = 'A3'

    wb.save(output_path)


# ══════════════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════════════

def main():
    print('='*70)
    print('  Word → Excel Question Bank Converter  (All Sections + Case-Based)')
    print('='*70)

    # Input folder
    while True:
        inp = input('\nEnter INPUT folder path (contains .docx files):\n> ').strip().strip('"\'')
        if not inp:               print('  ⚠  Path cannot be empty.');  continue
        if not os.path.isdir(inp):print(f'  ⚠  Folder not found: {inp}'); continue
        files = sorted(f for f in os.listdir(inp)
                       if f.lower().endswith('.docx') and not f.startswith('~$'))
        if not files: print(f'  ⚠  No .docx files in: {inp}'); continue
        break

    print(f'\n  Found {len(files)} file(s):')
    for f in files: print(f'    • {f}')

    # Output folder
    while True:
        out = input('\nEnter OUTPUT folder path:\n> ').strip().strip('"\'')
        if not out: print('  ⚠  Path cannot be empty.'); continue
        try:    os.makedirs(out, exist_ok=True)
        except Exception as e: print(f'  ⚠  Cannot create: {e}'); continue
        tst = os.path.join(out, '.wrtest')
        try:    open(tst,'w').close(); os.remove(tst)
        except: print(f'  ⚠  No write permission in: {out}'); continue
        break

    # Parse
    print('\n' + '='*70)
    print('  Parsing documents...')
    print('='*70)

    classes  = {}
    total_q  = 0
    ok_count = 0
    failed   = []

    for idx, fname in enumerate(files, 1):
        path = os.path.join(inp, fname)
        print(f'\n  [{idx}/{len(files)}]  {fname}')
        print(f'  {"─"*66}')
        try:
            cls, subj, rows = parse_document(path)
        except Exception as e:
            import traceback
            print(f'  ✖  ERROR: {e}')
            for ln in traceback.format_exc().splitlines()[-6:]:
                if ln.strip(): print(f'     {ln}')
            failed.append((fname, str(e))); continue

        if not rows:
            print('  ⚠  No questions found – skipped.')
            failed.append((fname, 'No questions found')); continue

        classes.setdefault(cls, []).extend(rows)

        by_type = {}
        for r in rows:
            qt = r['Question Type']
            by_type[qt] = by_type.get(qt,0)+1

        print(f'  ✅  {len(rows)} question(s) extracted')
        for qt in sorted(by_type): print(f'     {qt:<18} {by_type[qt]:>4}')
        total_q  += len(rows)
        ok_count += 1

    if total_q == 0:
        print('\n  ⚠  No questions extracted from any file.')
        input('\nPress Enter to exit...'); sys.exit(1)

    # Write Excel
    out_path = os.path.join(out, 'QuestionBank_Master.xlsx')
    print('\n' + '='*70)
    print(f'  Writing → {os.path.basename(out_path)}')
    print('='*70)

    if os.path.isfile(out_path):
        try: open(out_path,'a').close()
        except PermissionError:
            print('  ⚠  File is open in Excel – please close it.')
            input('     Press Enter to continue...')

    try:
        write_excel(classes, out_path)
    except PermissionError:
        print('\n  ✖  File locked in Excel. Close it and retry.')
        input('\nPress Enter to exit...'); sys.exit(1)
    except Exception as e:
        print(f'\n  ✖  Failed to write Excel: {e}')
        import traceback; traceback.print_exc()
        input('\nPress Enter to exit...'); sys.exit(1)

    # Summary
    print('\n  ✅  Excel created successfully!')
    print(f'  📁  {out_path}\n')
    for cn in sorted(classes):
        print(f'  Sheet: {sanitize_sheet(cn):<22} {len(classes[cn]):>5} question(s)')

    print()
    print('='*70)
    print(f'  Files processed   : {len(files)}')
    print(f'  ✅ Converted       : {ok_count}')
    print(f'  ✖  Failed/Skipped  : {len(failed)}')
    print(f'  Total questions   : {total_q}')
    if failed:
        print('\n  Failed files:')
        for fn,r in failed: print(f'    • {fn}  →  {r}')
    print('\n  💡 Tips:')
    print('     • Data → Filter to filter by Question Type, Chapter, etc.')
    print('     • Case-Based rows share the same Case Text (scenario)')
    print('     • Difficulty: L=Low  M=Mid  H=High  A=Advance')
    input('\nPress Enter to exit...')


if __name__ == '__main__':
    main()
